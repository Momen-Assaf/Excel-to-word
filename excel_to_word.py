import pandas as pd
from docx import Document
from docx.shared import Pt
import tkinter as tk
from tkinter import filedialog
import numpy as np

def excel_to_word(excel_path, word_path):
    # Read the Excel file
    excel_data = pd.read_excel(excel_path, sheet_name=None)  # Read all sheets

    # Create a new Word document
    doc = Document()

    # Set font to a common font that supports Arabic characters
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'  # You can change this to a different font if needed
    font.size = Pt(12)

    for sheet_name, sheet_data in excel_data.items():
        # Add a new section for each sheet
        doc.add_heading(sheet_name, level=1)

        # Replace NaN with empty strings
        sheet_data = sheet_data.replace(np.nan, '')

        # Add table to the document only if there is data to display
        if not sheet_data.empty:
            table = doc.add_table(rows=1, cols=sheet_data.shape[1])
            table.style = 'Table Grid'

            # Add header row
            hdr_cells = table.rows[0].cells
            for i, column in enumerate(sheet_data.columns):
                hdr_cells[i].text = str(column)

            # Add data rows, skipping empty rows
            for _, row in sheet_data.iterrows():
                if any(row):  # Skip the row if all fields are empty
                    row_cells = table.add_row().cells
                    for i, cell in enumerate(row):
                        if cell:  # Only add the cell if it's not empty
                            row_cells[i].text = str(cell)

    # Save the Word document
    doc.save(word_path)

def select_file():
    root = tk.Tk()
    root.withdraw()  # Hide the root window
    file_path = filedialog.askopenfilename(
        filetypes=[("Excel files", "*.xlsx *.xls")]
    )
    return file_path

def save_file():
    root = tk.Tk()
    root.withdraw()  # Hide the root window
    file_path = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Word files", "*.docx")]
    )
    return file_path

# Usage
excel_path = select_file()
if excel_path:
    word_path = save_file()
    if word_path:
        excel_to_word(excel_path, word_path)
        print(f"Conversion complete: {word_path}")
    else:
        print("Save operation cancelled.")
else:
    print("File selection cancelled.")
